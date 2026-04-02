import os
import tempfile
from typing import Dict, Any, List
from dataclasses import dataclass
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup
import re

@dataclass
class ProcessedDocument:
    """Document traité avec métadonnées."""
    text: str
    metadata: Dict[str, Any]
    sections: List[Dict[str, Any]]
    word_count: int
    has_tables: bool
    has_images: bool

class DocumentProcessor:
    """Processeur de documents pour extraire et structurer le texte."""
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.md': self._process_markdown
        }
    
    def process(self, file_content: bytes, file_extension: str) -> ProcessedDocument:
        """
        Traite un document et extrait le texte structuré.
        
        Args:
            file_content: Contenu brut du fichier
            file_extension: Extension du fichier (.pdf, .docx, etc.)
            
        Returns:
            ProcessedDocument: Document traité
        """
        file_extension = file_extension.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Extension non supportée: {file_extension}")
        
        # Créer un fichier temporaire pour le traitement
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        try:
            # Appeler le processeur approprié
            processor = self.supported_extensions[file_extension]
            result = processor(tmp_path)
            
            # Post-traitement commun
            result = self._post_process(result)
            
            return result
        finally:
            # Nettoyer le fichier temporaire
            os.unlink(tmp_path)
    
    def _process_pdf(self, file_path: str) -> ProcessedDocument:
        """Traite un fichier PDF."""
        text = ""
        metadata = {}
        has_tables = False
        has_images = False
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Métadonnées
            if pdf_reader.metadata:
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'pages': len(pdf_reader.pages)
                }
            
            # Extraire le texte de chaque page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"--- Page {page_num} ---\n{page_text}\n\n"
            
            # Détection simple de tables (basée sur les motifs)
            # Note: Pour une détection précise, utiliser camelot ou tabula
            table_patterns = [r'\|\s*[^|]+\s*\|', r'\+[-]+\+']
            for pattern in table_patterns:
                if re.search(pattern, text, re.MULTILINE):
                    has_tables = True
                    break
            
            # PDF peut contenir des images
            has_images = any('/XObject' in page.get('/Resources', {}) for page in pdf_reader.pages)
        
        return ProcessedDocument(
            text=text.strip(),
            metadata=metadata,
            sections=self._extract_sections(text),
            word_count=len(text.split()),
            has_tables=has_tables,
            has_images=has_images
        )
    
    def _process_docx(self, file_path: str) -> ProcessedDocument:
        """Traite un fichier DOCX."""
        doc = Document(file_path)
        text_parts = []
        metadata = {}
        has_tables = len(doc.tables) > 0
        has_images = any(paragraph._element.xpath('.//pic:pic') for paragraph in doc.paragraphs)
        
        # Métadonnées
        core_properties = doc.core_properties
        metadata = {
            'title': core_properties.title or '',
            'author': core_properties.author or '',
            'subject': core_properties.subject or '',
            'paragraphs': len(doc.paragraphs)
        }
        
        # Extraire le texte avec structure
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Conserver les styles (titre, sous-titre)
                style = paragraph.style.name if paragraph.style else 'Normal'
                text_parts.append(f"[{style}] {paragraph.text}")
        
        text = "\n".join(text_parts)
        
        return ProcessedDocument(
            text=text.strip(),
            metadata=metadata,
            sections=self._extract_sections(text),
            word_count=len(text.split()),
            has_tables=has_tables,
            has_images=has_images
        )
    
    def _process_text(self, file_path: str) -> ProcessedDocument:
        """Traite un fichier texte brut."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        return ProcessedDocument(
            text=text.strip(),
            metadata={'encoding': 'utf-8'},
            sections=self._extract_sections(text),
            word_count=len(text.split()),
            has_tables=False,
            has_images=False
        )
    
    def _process_markdown(self, file_path: str) -> ProcessedDocument:
        """Traite un fichier Markdown."""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Convertir en HTML puis extraire le texte
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        return ProcessedDocument(
            text=text.strip(),
            metadata={'format': 'markdown'},
            sections=self._extract_sections_from_markdown(md_content),
            word_count=len(text.split()),
            has_tables='|' in md_content,  # Détection simple de tables markdown
            has_images='![' in md_content  # Détection d'images markdown
        )
    
    def _extract_sections(self, text: str) -> List[Dict[str, Any]]:
        """Extrait les sections d'un texte basé sur les titres."""
        sections = []
        
        # Patterns pour détecter les titres
        title_patterns = [
            (r'^(#{1,6})\s+(.+)$', 'markdown'),  # Markdown headers
            (r'^([A-Z][A-Z\s]{5,})$', 'uppercase'),  # Titres en majuscules
            (r'^\d+\.\d+\s+.+$', 'numbered'),  # Numérotation 1.1, 2.3, etc.
        ]
        
        lines = text.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            for pattern, pattern_type in title_patterns:
                match = re.match(pattern, line)
                if match:
                    if current_section:
                        sections.append(current_section)
                    
                    title = match.group(2) if pattern_type == 'markdown' else match.group(1)
                    current_section = {
                        'title': title.strip(),
                        'type': pattern_type,
                        'level': len(match.group(1)) if pattern_type == 'markdown' else 1,
                        'start_line': i,
                        'content': []
                    }
                    break
            else:
                # Ce n'est pas un titre, ajouter au contenu de la section courante
                if current_section is not None:
                    current_section['content'].append(line)
                else:
                    # Texte avant le premier titre
                    if not sections or sections[-1]['title'] == 'Préambule':
                        if not sections:
                            sections.append({
                                'title': 'Préambule',
                                'type': 'preamble',
                                'level': 1,
                                'start_line': 0,
                                'content': [line]
                            })
                        else:
                            sections[-1]['content'].append(line)
        
        if current_section:
            sections.append(current_section)
        
        # Nettoyer les sections
        for section in sections:
            section['content'] = '\n'.join(section['content'])
            section['word_count'] = len(section['content'].split())
        
        return sections
    
    def _extract_sections_from_markdown(self, md_text: str) -> List[Dict[str, Any]]:
        """Extrait les sections spécifiquement depuis Markdown."""
        sections = []
        lines = md_text.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Détection des headers Markdown
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                if current_section:
                    sections.append(current_section)
                
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                current_section = {
                    'title': title,
                    'type': 'markdown',
                    'level': level,
                    'start_line': i,
                    'content': []
                }
            elif current_section is not None:
                current_section['content'].append(line)
            else:
                # Texte avant le premier header
                if not sections or sections[-1]['title'] == 'Préambule':
                    if not sections:
                        sections.append({
                            'title': 'Préambule',
                            'type': 'preamble',
                            'level': 1,
                            'start_line': 0,
                            'content': [line]
                        })
                    else:
                        sections[-1]['content'].append(line)
        
        if current_section:
            sections.append(current_section)
        
        # Nettoyer
        for section in sections:
            section['content'] = '\n'.join(section['content'])
            section['word_count'] = len(section['content'].split())
        
        return sections
    
    def _post_process(self, doc: ProcessedDocument) -> ProcessedDocument:
        """Post-traitement du document."""
        # Nettoyer le texte
        text = doc.text
        
        # Supprimer les espaces multiples
        text = re.sub(r'\s+', ' ', text)
        
        # Supprimer les sauts de ligne inutiles
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        
        # Normaliser les guillemets
        text = text.replace('"', '"').replace('"', '"')
        
        # Mettre à jour le document
        return ProcessedDocument(
            text=text.strip(),
            metadata=doc.metadata,
            sections=doc.sections,
            word_count=len(text.split()),
            has_tables=doc.has_tables,
            has_images=doc.has_images
        )